# -*- coding: utf-8 -*-
"""中文 DOCX 生成助手（python-docx）。
提供统一的中文字体、标题、勾选项、着色表格等函数，用于生成
"审查要点" 与 "审查意见书" 等专业文档，保持排版一致。

依赖：python-docx
  venv: /Users/bob/.workbuddy/binaries/python/envs/default/bin/python
  安装: .../bin/pip install python-docx

典型用法见本 skill 的 SKILL.md。核心 API：
  d = Doc(out_path)
  d.title("标题", color=(0xC0,0,0))
  d.h("一、总则", 1)
  d.p("正文…", color=(0x59,0x59,0x59))
  d.check("检查项文字", tag="5.1.1c")       # 生成 "☐ … ［5.1.1c］"
  d.table(header=[...], rows=[[...],...], header_fill="1F4E79")
  d.save()
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def _rgb(c):
    if c is None:
        return None
    if isinstance(c, RGBColor):
        return c
    return RGBColor(c[0], c[1], c[2])


def set_cn_font(run, name="宋体", size=None, bold=None, color=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), name)
    rFonts.set(qn('w:hAnsi'), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = _rgb(color)


def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hexcolor)
    tcPr.append(shd)


class Doc:
    def __init__(self, out_path):
        self.out = out_path
        self.doc = Document()
        style = self.doc.styles['Normal']
        style.font.name = '宋体'; style.font.size = Pt(10.5)
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        for s in self.doc.sections:
            s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
            s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

    def title(self, text, size=20, color=(0xC0, 0x00, 0x00)):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); set_cn_font(r, "黑体", size, True, color)
        return p

    def subtitle(self, text, size=10.5, color=(0x59, 0x59, 0x59)):
        p = self.doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text); set_cn_font(r, "宋体", size, False, color)
        return p

    def h(self, text, level=1):
        colors = {1: (0x1F, 0x4E, 0x79), 2: (0x2E, 0x5C, 0x8A), 3: (0x3C, 0x6E, 0x2E)}
        sizes = {1: 15, 2: 13, 3: 11.5}
        head = self.doc.add_heading(level=level); r = head.add_run(text)
        set_cn_font(r, "黑体", sizes.get(level, 10.5), True, colors.get(level, (0, 0, 0)))
        return head

    def p(self, text="", size=10.5, bold=False, color=None, align=None, italic=False):
        para = self.doc.add_paragraph(); r = para.add_run(text)
        set_cn_font(r, "宋体", size, bold, color)
        if italic:
            r.font.italic = True
        if align is not None:
            para.alignment = align
        return para

    def check(self, text, tag=None, level=0):
        box = "☐ " if level == 0 else ("    ☐ " if level == 1 else "        ☐ ")
        suffix = ("  ［%s］" % tag) if tag else ""
        para = self.doc.add_paragraph(box + text + suffix)
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.space_before = Pt(1)
        for r in para.runs:
            set_cn_font(r, "宋体", 10.5)
            if tag and ("［" in r.text):
                r.font.color.rgb = _rgb((0x8B, 0x45, 0x00)); r.font.size = Pt(9)
        return para

    def bullet(self, text, sign="•"):
        para = self.doc.add_paragraph("  " + sign + " " + text)
        para.paragraph_format.space_after = Pt(1)
        for r in para.runs:
            set_cn_font(r, "宋体", 10.5)
        return para

    def table(self, header, rows, header_fill="1F4E79", header_color=(0xFF, 0xFF, 0xFF),
              size=9.5, widths=None, red_last_col=False):
        t = self.doc.add_table(rows=1, cols=len(header))
        t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
        hc = t.rows[0].cells
        for i, txt in enumerate(header):
            self._cell(hc[i], txt, True, size, header_color, header_fill)
        for row in rows:
            c = t.add_row().cells
            for j, txt in enumerate(row):
                col = (0xC0, 0x00, 0x00) if (red_last_col and j == len(row) - 1) else None
                self._cell(c[j], txt, False, size, col)
        if widths:
            for r in t.rows:
                for j, w in enumerate(widths):
                    r.cells[j].width = Cm(w)
        return t

    def _cell(self, cell, text, bold=False, size=9.5, color=None, fill=None):
        cell.text = ""; para = cell.paragraphs[0]; r = para.add_run(text)
        set_cn_font(r, "宋体", size, bold, color)
        if fill is not None:
            shade_cell(cell, fill)

    def spacer(self):
        self.doc.add_paragraph()

    def save(self):
        self.doc.save(self.out)
        return self.out
