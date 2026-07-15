# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/bob/Desktop/T2/220kV常规站继电保护二次安措票审查要点.docx"
doc = Document()

def set_cn_font(run, name="宋体", size=None, bold=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name); rFonts.set(qn('w:ascii'), name); rFonts.set(qn('w:hAnsi'), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold

style = doc.styles['Normal']
style.font.name = '宋体'; style.font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0); s.left_margin = Cm(2.2); s.right_margin = Cm(2.2)

def add_heading(text, level=1):
    h = doc.add_heading(level=level); run = h.add_run(text)
    if level == 1:
        set_cn_font(run, "黑体", 15, True); run.font.color.rgb = RGBColor(0x1F,0x4E,0x79)
    elif level == 2:
        set_cn_font(run, "黑体", 13, True); run.font.color.rgb = RGBColor(0x2E,0x5C,0x8A)
    elif level == 3:
        set_cn_font(run, "黑体", 11.5, True); run.font.color.rgb = RGBColor(0x3C,0x6E,0x2E)
    else:
        set_cn_font(run, "黑体", 10.5, True)
    return h

def add_para(text, size=10.5, bold=False, color=None, align=None, italic=False):
    p = doc.add_paragraph(); run = p.add_run(text)
    set_cn_font(run, "宋体", size, bold)
    if color is not None: run.font.color.rgb = color
    if align is not None: p.alignment = align
    if italic: run.font.italic = True
    return p

def add_check(text, tag=None, level=0):
    box = "☐ " if level == 0 else ("    ☐ " if level == 1 else "        ☐ ")
    suffix = ("  ［%s］" % tag) if tag else ""
    p = doc.add_paragraph(box + text + suffix)
    p.paragraph_format.space_after = Pt(1); p.paragraph_format.space_before = Pt(1)
    for r in p.runs:
        set_cn_font(r, "宋体", 10.5)
        if tag and r.text.endswith("］") and ("%" not in r.text):
            pass
    # 条款号着色
    if tag:
        for r in p.runs:
            if r.text.startswith("  ［") or r.text.startswith("［"):
                r.font.color.rgb = RGBColor(0x8B,0x45,0x00); r.font.size = Pt(9)
    return p

def add_bullet(text, sign="•"):
    p = doc.add_paragraph("  " + sign + " " + text)
    p.paragraph_format.space_after = Pt(1)
    for r in p.runs: set_cn_font(r, "宋体", 10.5)
    return p

def shade_cell(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexcolor); tcPr.append(shd)

def set_cell(cell, text, bold=False, size=10, color=None, fill=None):
    cell.text = ""; p = cell.paragraphs[0]; run = p.add_run(text)
    set_cn_font(run, "宋体", size, bold)
    if color is not None: run.font.color.rgb = color
    if fill is not None: shade_cell(cell, fill)

# ===== 封面 =====
t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("220kV常规站继电保护二次安措票审查要点")
set_cn_font(r, "黑体", 20, True); r.font.color.rgb = RGBColor(0xC0,0x00,0x00)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("（依据《220kV常规变电站继电保护安全措施票技术细则》第5章 编制｜覆盖保护检验与保护改造）")
set_cn_font(r, "宋体", 10.5, False); r.font.color.rgb = RGBColor(0x59,0x59,0x59)
doc.add_paragraph()
note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = note.add_run("适用：220kV常规站 主变 / 母差 / 线路保护 检验与改造（更换）二次安措票")
set_cn_font(r, "宋体", 9.5, False); r.font.color.rgb = RGBColor(0x80,0x80,0x80)
doc.add_paragraph()

# ===== 一、总则 =====
add_heading("一、总则与适用范围", 1)
add_para("本审查要点依据《220kV常规变电站继电保护安全措施票技术细则》第5章（5.1 保护检验、5.2 保护改造）编制，用于继电保护二次安措票在编制、审核、签发及执行前的审查。常规站以硬压板投退 + 端子排解线/封挡为核心手段，所有解线须“用红色绝缘胶布包好”、断开空开须“用红色绝缘胶布封挡”。")
add_para("要点按“原始状态确认 → 联跳回路 → 采样回路（电流/电压）→ 信号回路 → 交直流电源 → 过程安全”逐环节列出，每条标注对应细则条款号（如 5.1.1c），便于溯源。", color=RGBColor(0x59,0x59,0x59))
add_heading("审查依据（核对时必备）", 3)
add_bullet("《220kV常规变电站继电保护安全措施票技术细则》第5章及附录 A.1/A.2/A.3")
add_bullet("本站二次电缆信息流图（含压板标识）——核对端子号、回路标号、压板命名")
add_bullet("保护装置出厂/施工图纸、端子排图")
add_bullet("调度命名、设备双重名称、工作票及停电范围")

# ===== 二、审查流程 =====
add_heading("二、审查组织与流程", 1)
for k, v in [
    ("① 编制自审", "编制人按模板填写后，逐项对照本要点自查，重点核对联跳、电流、电压回路完整性与端子编号，并标注作业类型（检验/改造）。"),
    ("② 技术审核", "班组长或技术负责人复核回路逻辑、与信息流图一致性、过程安措可逆性。"),
    ("③ 签发", "工作票签发人（或继电保护专业负责人）最终审查并签字；补充安措须经签发人确认。"),
    ("④ 执行前复核", "开工前工作负责人与监护人依据原始状态记录（含定值区号、拍照）核对现场实际压板/空开状态，确认无误后执行。"),
]:
    p = doc.add_paragraph(); r1 = p.add_run(k+"："); set_cn_font(r1,"黑体",10.5,True)
    r2 = p.add_run(v); set_cn_font(r2,"宋体",10.5); p.paragraph_format.space_after = Pt(2)

# ===== 三、票面完整性 =====
add_heading("三、票面完整性审查（每张票必查）", 1)
add_check("被试设备名称与调度命名、双重编号、工作票一致")
add_check("工作内容明确，并注明作业类型（保护检验 / 保护改造（含更换））")
add_check("工作负责人、工作时间（起止）、签发人栏填写完整且具备相应资质")
add_check("原始状态记录表完整：列出开工前所有相关屏柜硬压板、直流/交流空开、切换把手、光纤芯号状态，设“开工确认”与“收工确认”两栏对应")
add_check("安全措施表按“执行 / 安全措施内容 / 恢复”三栏填写，与原始状态一一对应、可逆向恢复")
add_check("端子号、回路标号、压板命名与信息流图/图纸完全一致，无“XX（）”等模糊待填项")
add_check("补充安措说明栏完整，并经签发人确认")
add_check("签字栏齐全：执行人、执行监护人、恢复人、恢复监护人")
add_check("绝缘处理规范：解线/解端子“用红色绝缘胶布包好”；断开空开“用红色绝缘胶布封挡”；光纤“加装防尘帽/防尘套”")

# ===== 四、原始状态确认 =====
add_heading("四、原始状态确认审查（5.1.1 / 5.2.1）", 1)
add_heading("通用项", 3)
add_check("记录保护装置定值区号", "5.1.1a")
add_check("拍照记录保护屏前压板、屏后空开、切换把手、光纤芯号", "5.1.1b")
add_heading("保护检验（5.1.1 c/d/e）", 3)
add_check("线路保护检验：确认母差保护上该支路跳闸出口及启动失灵压板已退出，并封挡压板", "5.1.1c")
add_check("主变保护检验：确认各侧母差保护屏上该支路跳闸出口、启动失灵、解复压闭锁压板已退出；主变保护屏上联跳母联、分段及其他并网电源开关出口压板已退出，并封挡", "5.1.1d")
add_check("母差保护检验：确认母差屏上跳各运行支路开关的跳闸及主变失灵联跳出口压板、启动失灵开入压板已退出，并封挡", "5.1.1e")
add_heading("保护改造（5.2.1 a/b/c/d）", 3)
add_check("线路保护改造：确认母差保护上该支路跳闸出口及启动失灵压板已退出；测控屏对应遥控电源、遥信电源已断开并封挡", "5.2.1a")
add_check("主变保护改造：确认各侧母差保护屏上该支路跳闸出口、启动失灵、解复压闭锁压板已退出；测控屏对应遥控、遥信电源空开已断开并封挡", "5.2.1b")
add_check("母差保护改造：确认运行于该母线的主变保护屏上失灵联跳开入压板已退出并封挡", "5.2.1c")
add_check("保护装置改造：确认直流分电柜上装置电源、操作电源等相关空开已断开并封挡", "5.2.1d")

# ===== 五、联跳回路 =====
add_heading("五、联跳回路安全措施审查（5.1.2 / 5.2.2）  ［最高风险·优先审查］", 1)
add_para("联跳/失灵回路漏退、漏解将直接误跳运行设备，是审查的重中之重。", color=RGBColor(0xC0,0x00,0x00))
add_heading("保护检验（5.1.2 a/b/c）", 3)
add_check("线路保护检验：拔下保护通道光纤，并加装防尘帽", "5.1.2a")
add_check("主变保护检验：解除主变保护屏上联跳母联、分段开关及其他并网电源开关的接线端子，并封挡", "5.1.2b")
add_check("母差保护检验：解除连接至各间隔保护屏柜的跳闸回路二次线，以及主变保护屏柜的失灵联跳、解除复压闭锁等接线端子，并封挡", "5.1.2c")
add_heading("保护改造（5.2.2 a/b/c）", 3)
add_check("线路保护改造：解除母差保护至该支路启动失灵回路及跳闸回路接线端子，并封挡", "5.2.2a")
add_check("主变保护改造：解除母差保护至该支路启动失灵、失灵联跳及解复压闭锁回路接线端子，并封挡", "5.2.2b")
add_check("母差保护改造：解除所有跳闸出口回路、启动失灵回路、主变解除复压闭锁及失灵联跳回路接线端子，并封挡（不得遗漏任一运行间隔）", "5.2.2c")
add_heading("通用核对", 3)
add_check("解除端子标注屏柜、端子号、回路含义（与信息流图一致），并“用红色绝缘胶布包好”")
add_check("同一回路“解除—投入”成对、对称：执行列退出/解除，恢复列对应恢复")

# ===== 六、采样回路 =====
add_heading("六、采样回路安全措施审查（5.1.3 / 5.2.3）  ［防 CT 开路 / PT 短路］", 1)
add_heading("电流回路", 2)
add_check("一般要求：确认装置采样无电流后，在保护屏外侧短接并打开电流试验端子连片，并封挡端子排", "5.1.3a1")
add_check("母差检验/改造：先在母差屏电流端子外侧短接运行支路电流二次回路，确认各运行支路采样均无电流，再打开电流试验端子连片，封挡", "5.1.3a2")
add_check("串接装置（故障录波器、行波测距、安全自动装置等）：保护装置电流回路尾部短接，在串接装置确认无电流后，打开至上述装置的电流试验端子连片，并封挡", "5.1.3a3")
add_check("CT 通流等试验：二次电流回路须物理隔离，在端子箱内至运行设备侧短接电流回路，确认对应装置无电流后打开至运行设备电流回路试验端子连片，并封挡", "5.1.3a4")
add_check("改造（线路/主变）：确认无电流后打开端子箱母差电流端子连接片并封挡，同时在母差屏后将该间隔电流端子连接片打开并封挡（两处均处理）", "5.2.3a1/a2")
add_check("母差套屏改造：各支路端子箱 CT 侧短接母差用电流绕组二次回路，确认无电流后打开连片封挡；同时在母差屏后将该间隔电流端子连接片打开封挡", "5.2.3a3")
add_check("母差立新屏改造：支路端子箱 CT 侧短接至母差电流二次回路，确认无电流打开连片，解除支路端子箱至原母差屏电流回路接线端子并封挡", "5.2.3a4")
add_check("所有相及中性线（A/B/C 及 AN/BN/CN）端子均封挡，无遗漏；仅处理本工作间隔，严禁误碰运行间隔")
add_heading("电压回路", 2)
add_check("打开电压试验端子连片或解除保护屏内电压接线端子，并封挡", "5.1.3b1/5.2.3b")
add_check("采用线路电压互感器：确认开关端子箱内线路 PT 二次空开已断开，并封挡", "5.1.3b2")
add_check("电压中性线 N600 已一并处理，解线后绝缘包扎，防止相间/对地短路")

# ===== 七、信号回路 =====
add_heading("七、信号回路安全措施审查（5.2.4）", 1)
add_check("解除故障录波器至保护屏开关量回路的公共端及相关开关量回路接线端子，并封挡", "5.2.4a")
add_check("解除测控装置至保护屏信号回路的公共端及相关信号回路接线端子，并封挡", "5.2.4b")

# ===== 八、交直流电源 =====
add_heading("八、交直流电源回路安全措施审查（5.2.5 / 5.2.1d）", 1)
add_check("直流：解除直流分电屏至本屏直流电源的接线端子并封挡；装置电源、操作电源等相关空开已断开并封挡", "5.2.5a/5.2.1d")
add_check("交流：明确交流电源来源及去向，解除交流电源来源侧接线端子，并封挡", "5.2.5b")
add_check("空开/端子编号与图纸一致，装置电源、操作电源、交换机电源区分清楚，严禁误断运行设备电源")

# ===== 九、过程安全 =====
add_heading("九、过程安全措施审查（5.2.6）  ［临时·可逆］", 1)
add_para("过程安措为试验期间临时投退，审查重点是“试验前退出、试验后恢复、逻辑闭环”。", color=RGBColor(0x59,0x59,0x59))
add_check("线路及主变保护改造接入母差传动试验：确认母差保护上所有硬压板已退出，并封挡压板", "5.2.6a")
add_check("母差保护改造某间隔接入：确认母差所有硬压板均已退出，并恢复该间隔电流回路安全措施；接入完成后封挡端子排", "5.2.6b")
add_check("试验仅投入本间隔相关压板，试验结束后明确恢复项，逻辑闭环")

# ===== 十、三类典型工作专项速查 =====
add_heading("十、三类典型工作专项速查", 1)
add_para("下表按设备类型汇总最关键的审查项与对应条款，便于快速定位（详细检查项见第四~九章）。", color=RGBColor(0x59,0x59,0x59))
tbl = doc.add_table(rows=1, cols=4); tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl.rows[0].cells
for i, txt in enumerate(["作业类型","原始状态关键项（条款）","联跳/采样关键项（条款）","过程安措（条款）"]):
    set_cell(hdr[i], txt, True, 9.5, fill="1F4E79")
    for p in hdr[i].paragraphs:
        for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
rows = [
    ("主变保护\n（检验/改造）","各侧母差：跳闸出口、启动失灵、解复压闭锁压板退出封挡；联跳母联/分段/并网电源出口退出封挡（5.1.1d/5.2.1b）","解除母差至本支路启动失灵/失灵联跳/解复压闭锁端子（5.1.2b/5.2.2b）；各侧电流端子箱+屏后双点封挡（5.2.3a2）；多电压等级电压（5.2.3b）","接入两套母差传动：母差所有硬压板退出封挡（5.2.6a）"),
    ("母差保护\n（检验/改造）","运行于该母线主变失灵联跳开入压板退出封挡（5.1.1e/5.2.1c）","解除至全部间隔（母联/各线路/各主变）跳闸及启动失灵、主变解复压闭锁/失灵联跳端子（5.1.2c/5.2.2c）；各运行支路电流先短接确认无流再打开（5.1.3a2/5.2.3a3）","某间隔接入：母差硬压板全退、恢复该间隔电流回路，完成封挡（5.2.6b）"),
    ("线路保护\n（检验/改造）","母差上该支路跳闸出口、启动失灵压板退出封挡；测控遥控/遥信电源断开封挡（5.1.1c/5.2.1a）","检验拔通道光纤装防尘帽（5.1.2a）；改造解除母差至本支路启动失灵+跳闸端子（5.2.2a）；本线路电流端子箱+屏后双点封挡（5.2.3a1）；线路PT二次空开断开（5.1.3b2）","接入两套母差传动：母差所有硬压板退出封挡（5.2.6a）"),
]
for row in rows:
    cells = tbl.add_row().cells
    for j, txt in enumerate(row):
        set_cell(cells[j], txt, False, 9)
for r in tbl.rows:
    r.cells[0].width = Cm(2.6); r.cells[1].width = Cm(5.2); r.cells[2].width = Cm(5.6); r.cells[3].width = Cm(3.6)

# ===== 十一、高危红线 =====
add_heading("十一、高危红线与典型错误清单", 1)
add_para("以下任一项未通过，安措票不得签发、不得执行。", color=RGBColor(0xC0,0x00,0x00), bold=True)
tbl2 = doc.add_table(rows=1, cols=3); tbl2.style = 'Table Grid'; tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = tbl2.rows[0].cells
for i, txt in enumerate(["序号","高危红线 / 典型错误","后果"]):
    set_cell(hdr[i], txt, True, 10, fill="1F4E79")
    for p in hdr[i].paragraphs:
        for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
redlines = [
    ("1","联跳/失灵回路漏退、漏解（遗漏对侧或相邻间隔）","误跳运行设备 ⚠"),
    ("2","电流回路未先短接即打开连片，或漏封中性线","CT 二次开路，危险高电压 ⚠"),
    ("3","电压回路解线未绝缘包扎 / 仅解线未打开试验端子连片","PT 二次短路/接地，烧毁设备 ⚠"),
    ("4","红色绝缘胶布漏包/漏封挡；光纤未装防尘帽/防尘套","误碰带电端子、插错光纤"),
    ("5","母差改造漏解任一运行间隔（线路/主变/母联）跳闸回路","误跳运行设备 ⚠"),
    ("6","改造电流仅处理单点（漏端子箱或漏屏后），隔离不彻底","运行设备误动/开路"),
    ("7","串接装置（录波/行波/安自）电流未尾部短接即打开连片","误动、CT 开路 ⚠"),
    ("8","母差立新屏改造未短接至新屏电流回路即解原屏接线","运行设备误动"),
    ("9","CT 通流试验未做物理隔离、端子箱侧未先短接","CT 开路/误动"),
    ("10","过程安措母差硬压板未全退 / 未恢复该间隔电流回路","误动、状态不一致"),
    ("11","空开误断运行设备电源（装置/操作/交换机）","运行设备失电、保护退出"),
    ("12","未记录定值区号、未拍照（压板/空开/把手/光纤芯）","原始状态无据，收工难核对"),
    ("13","电压多套/多电压等级回路漏项（220/110/35kV，一/二套）","运行保护失压、误动"),
    ("14","作业类型（检验/改造）填错导致措施不匹配；签字不全或未经签发确认","流程不合规、责任不清"),
]
for row in redlines:
    cells = tbl2.add_row().cells
    set_cell(cells[0], row[0], True, 9.5)
    set_cell(cells[1], row[1], False, 9.5)
    set_cell(cells[2], row[2], False, 9.5, color=RGBColor(0xC0,0x00,0x00))
for r in tbl2.rows:
    r.cells[0].width = Cm(1.2); r.cells[1].width = Cm(11.0); r.cells[2].width = Cm(4.0)

# ===== 十二、审查结论 =====
add_heading("十二、审查结论与签字", 1)
add_para("审查结论：☐ 合格，同意签发    ☐ 修改后复审    ☐ 不合格", size=11, bold=True)
add_para("")
ct = doc.add_table(rows=4, cols=4); ct.style = 'Table Grid'
for i, row in enumerate([("审查项目","审查人","审查日期","结论"),("编制自审","","",""),("技术负责人审核","","",""),("签发人签发","","","")]):
    cells = ct.rows[i].cells
    if i == 0:
        for j, txt in enumerate(row): set_cell(cells[j], txt, True, 10, fill="D9E2F3")
    else:
        for j, txt in enumerate(row): set_cell(cells[j], txt, False, 10)
add_para("")
add_para("说明：本审查要点依据《220kV常规变电站继电保护安全措施票技术细则》第5章（5.1 保护检验、5.2 保护改造）及本站主变/母差/线路保护安措票模板提炼，供现场审核参照；具体项目以本站图纸、信息流图及调度要求为准。",
         size=9, color=RGBColor(0x80,0x80,0x80), italic=True)

doc.save(OUT)
print("Saved:", OUT)
