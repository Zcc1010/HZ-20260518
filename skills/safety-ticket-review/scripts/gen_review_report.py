# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = "/Users/bob/Desktop/T2/菊江变2号主变保护更换安全措施票_审查意见.docx"
doc = Document()

def set_cn_font(run, name="宋体", size=None, bold=None):
    run.font.name = name
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name); rFonts.set(qn('w:ascii'), name); rFonts.set(qn('w:hAnsi'), name)
    if size is not None: run.font.size = Pt(size)
    if bold is not None: run.font.bold = bold

style = doc.styles['Normal']; style.font.name='宋体'; style.font.size=Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
for s in doc.sections:
    s.top_margin=Cm(2.0); s.bottom_margin=Cm(2.0); s.left_margin=Cm(2.2); s.right_margin=Cm(2.2)

def H(text, level=1):
    h=doc.add_heading(level=level); r=h.add_run(text)
    if level==1: set_cn_font(r,"黑体",15,True); r.font.color.rgb=RGBColor(0x1F,0x4E,0x79)
    elif level==2: set_cn_font(r,"黑体",13,True); r.font.color.rgb=RGBColor(0x2E,0x5C,0x8A)
    elif level==3: set_cn_font(r,"黑体",11.5,True); r.font.color.rgb=RGBColor(0x3C,0x6E,0x2E)
    return h

def P(text, size=10.5, bold=False, color=None, align=None):
    p=doc.add_paragraph(); r=p.add_run(text); set_cn_font(r,"宋体",size,bold)
    if color is not None: r.font.color.rgb=color
    if align is not None: p.alignment=align
    return p

def shade(cell, hexcolor):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hexcolor); tcPr.append(shd)

def cell(c, text, bold=False, size=9.5, color=None, fill=None):
    c.text=""; p=c.paragraphs[0]; r=p.add_run(text); set_cn_font(r,"宋体",size,bold)
    if color is not None: r.font.color.rgb=color
    if fill is not None: shade(c,fill)

# ===== 封面 =====
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("继电保护二次安措票\n审 查 意 见 书"); set_cn_font(r,"黑体",20,True); r.font.color.rgb=RGBColor(0xC0,0x00,0x00)
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sub.add_run("被审文件：菊江变2号主变保护更换安全措施票（常规站）"); set_cn_font(r,"宋体",11,False); r.font.color.rgb=RGBColor(0x59,0x59,0x59)
doc.add_paragraph()

meta=[("被审文件","菊江变2号主变保护更换安全措施票.docx"),
      ("设备/作业类型","220kV菊江变 2号主变保护装置更换（常规站·主变保护改造）"),
      ("审查依据","《220kV常规变电站继电保护安全措施票技术细则》第5章（5.1检验/5.2改造）及《常规站二次安措票审查要点》"),
      ("审查结论","修改后复审（存在1项高危电压端子错误、1项电流中性线标错及多项封挡/电源/过程安措缺失）")]
mt=doc.add_table(rows=len(meta),cols=2); mt.style='Table Grid'
for i,(k,v) in enumerate(meta):
    cell(mt.rows[i].cells[0],k,True,10,fill="D9E2F3")
    cell(mt.rows[i].cells[1],v,False,10)
    mt.rows[i].cells[0].width=Cm(3.2); mt.rows[i].cells[1].width=Cm(13.5)

# ===== 统计 =====
P("")
P("问题统计：必须整改 9 项 ｜ 建议核实 7 项 ｜ 其中高危（可能误拆/误动运行设备）2 项。", bold=True, color=RGBColor(0xC0,0x00,0x00))

# ===== 一、必须整改 =====
H("一、必须整改项（签发前须闭环）",1)
must=[["M1","票面·编号/签发人/工作负责人","编号栏空白；签发人栏空白；表头缺「工作负责人」字段（仅被试设备/工作时间/签发人）。","第三章 票面完整性","补填编号、签发人，并增设工作负责人栏及签字。"],
["M2","原始状态 5.2.1b","缺各侧母差屏跳对应电压等级开关出口压板确认。5.2.1b要求各侧母差该支路跳闸出口压板已退出：220kV母差（双套A/B屏）跳2802开关、110kV母差（单套）跳402开关、35kV母差（若配置）跳302开关；本票仅确认启动失灵、解复压闭锁，未涉及各侧跳闸出口压板。","5.2.1b","补充确认220kV母差A/B屏跳2802开关、110kV母差屏跳402开关、35kV母差屏（若配置）跳302开关的出口压板已退出并封挡（主变全停可保留须记载说明）。"],
["M3","原始状态/电源 5.2.1b","缺测控屏遥控电源空开断开与封挡。5.2.1b要求遥控、遥信电源均断开封挡；本票仅含遥信电源（4xDKYX），无遥控电源。","5.2.1b","补充遥控电源空开的记录、断开与封挡。"],
["M4","采样·电流 D1","35kV电流中性线端子标错：302开关柜1ID18标为A330（应为N330）。","5.2.3a2/红线②","更正为N330；中性线须正确识别并封挡，避免误判相别/开路。"],
["M5","采样·电压 D3【高危】","110kVII母零序（A屏）条目终端号与标签错误，与220kVII母零序(A屏)完全相同：1-7D-40 L640/2BYH-131，系复制220kV条目所致；实际应为2-7D-40 L640/2YYH-132。该错误会误拆220kV电压端子、且110kV电压实际未隔离。","5.2.3b/红线③","必须更正终端与标签，执行前与图纸、信息流图逐字核对。"],
["M6","交直流 5.2.5","直流空开、遥信电源空开断开后均未标注用红色绝缘胶布封挡。","5.2.5a/5.2.1d","断开后统一标注红胶布封挡，防误合。"],
["M7","交直流 5.2.5b","缺交流电源回路安全措施。5.2.5b要求明确交流电源来源去向并解除来源侧接线端子封挡；本票电源回路仅含直流。","5.2.5b","补充相邻屏至本屏交流电源L/N的接线解除与封挡。"],
["M8","过程安措 5.2.6","缺少六、过程安全措施章节。主变保护更换含接入母差传动试验时须确认母差所有硬压板退出并封挡（5.2.6a）。","5.2.6","若工程含传动试验，必须补充过程安措章节及临时投退/恢复项。"],
["M9","票面·回路分类格式","回路分类顺序与模板不符（模板顺序：联跳→电流→电压→信号→交直流电源→过程；本票顺序：电源→信号→启失灵及解复压→电流→电压），分类名称与模板不一致（如「启失灵及解复压」应为「联跳回路」、「电源回路」应为「交直流电源回路」），各分组内同类型措施缺小标题区分（如启失灵回路下A屏/B屏措施应分别加小标题），且各组间无空行分隔。","四、票面规范性","应按模板标准调整为：一、联跳回路安全措施 → 二、电流回路安全措施 → 三、电压回路安全措施 → 四、信号回路安全措施 → 五、交直流电源回路安全措施 → 六、过程安全措施；各组间添加空行；各分组内同类型措施加小标题（格式：具体回路描述+冒号，加粗）；改造票须包含六、过程安全措施。"]]
tbl=doc.add_table(rows=1,cols=5); tbl.style='Table Grid'; tbl.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr=["编号","位置/条款","问题描述","依据","整改建议"]
hc=tbl.rows[0].cells
for i,txt in enumerate(hdr):
    cell(hc[i],txt,True,9.5,fill="1F4E79")
    for p in hc[i].paragraphs:
        for rr in p.runs: rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
for row in must:
    c=tbl.add_row().cells
    for j,txt in enumerate(row[:5]):
        cell(c[j],txt,False,9)
    c[0].width=Cm(1.0); c[1].width=Cm(2.6); c[2].width=Cm(6.0); c[3].width=Cm(2.0); c[4].width=Cm(4.0)

# ===== 二、建议核实 =====
doc.add_paragraph()
H("二、建议核实项（须对照本站图纸/信息流图确认）",1)
adv=[["S1","联跳 5.2.2b","缺【母差→主变失灵联跳回路】解除。5.2.2b要求解除母差至本支路启动失灵、失灵联跳及解复压闭锁；本票仅解除启动失灵+解复压闭锁。","核实本站母差失灵联跳是否经主变保护屏；若是须补充解除并封挡。"],
["S2","联跳 端子标号","B屏启动失灵负端【1C3D18 2B-138A/SL103'】、解复压负端【1C3D19 2B-138A/JFY103'】，负端出现SL103'/JFY103'（常规应为SL102/JFY102）。","与图纸核对正负端子标号一致性。"],
["S3","采样·电流 5.2.3a2","电流回路仅处理端子箱侧（短接+打开中连片），缺【母差屏后该间隔电流端子连接片打开并封挡】双点处理。","核实本站做法；若按细则须补屏后侧处理。"],
["S4","信号 5.2.4b","缺【测控装置至保护屏信号回路公共端】解除。本票仅断开遥信电源空开，未解除测控→保护屏信号回路接线。","核实并补充解除测控至保护屏信号回路公共端及接线。"],
["S5","采样·电压 35kV","35kV电压仅拆至A屏、标注【35kVI段母线压变】却接【35kVII母】，B屏未见。","核对35kV电压实际接线（I段/II母一致性及B屏是否需处理）。"],
["S6","被试设备名称","名称含【2号主变保护测控装置】，但工作内容仅为【保护装置更换】（测控未更换）。","核对测控是否在本次范围，避免歧义。"],
["S7","原始状态·定值/拍照","本票为改造（5.2.1未强制），但建议按良好实践补充定值区号记录与压板/空开/把手/光纤芯号拍照。","建议补充，便于收工核对与追溯。"]]
tbl2=doc.add_table(rows=1,cols=4); tbl2.style='Table Grid'; tbl2.alignment=WD_TABLE_ALIGNMENT.CENTER
hc=tbl2.rows[0].cells
for i,txt in enumerate(["编号","位置/条款","问题描述","核实/建议"]):
    cell(hc[i],txt,True,9.5,fill="2E5C8A")
    for p in hc[i].paragraphs:
        for rr in p.runs: rr.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
for row in adv:
    c=tbl2.add_row().cells
    for j,txt in enumerate(row):
        cell(c[j],txt,False,9)
    c[0].width=Cm(1.0); c[1].width=Cm(2.8); c[2].width=Cm(7.5); c[3].width=Cm(5.5)

# ===== 三、审查明细（按章节） =====
doc.add_paragraph()
H("三、逐章审查明细",1)
H("（一）票面完整性（第三章）",3)
P("■ 问题：编号、签发人、工作负责人缺失（M1）。")
P("✓ 被试设备名称、工作内容（主变保护装置更换）明确，属改造类，作业类型标注一致。")
H("（二）票面规范性（第四章）",3)
P("■ 问题：回路分类顺序与模板不符、分类名称不一致、缺小标题区分、无空行分隔（M9）。")
P("正确应为：一、联跳回路安全措施 → 二、电流回路安全措施 → 三、电压回路安全措施 → 四、信号回路安全措施 → 五、交直流电源回路安全措施 → 六、过程安全措施（改造票须有）；各分组内同类型措施加小标题（如「220kV 1号变压器第一套保护启动220kV第一套母差保护失灵及解复压闭锁回路：」）。")
P("✓ 各分类标题加粗标注。")
H("（三）原始状态确认（第五章 / 5.2.1）",3)
P("■ 问题：缺母差跳主变各侧开关出口压板确认（M2）；缺遥控电源（M3）；定值区号/拍照建议补充（S7）。")
P("✓ 启动失灵、解复压闭锁压板确认退出；直流/遥信电源空开已记录。")
H("（四）联跳回路（第六章 / 5.2.2b）",3)
P("■ 问题：缺母差→主变失灵联跳回路（S1）；B屏负端子标号异常（S2）。")
P("✓ 启动失灵（正/负）、解复压闭锁（正/负）均按双极性拆除并红胶布包好，方向正确。")
H("（五）采样回路（第七章 / 5.2.3）",3)
P("■ 问题：35kV电流中性线标错A330（M4）；110kVII母零序电压终端/标签复制错误（M5，高危）；电流仅单点处理（S3）；35kV电压范围/屏位待核（S5）。")
P("✓ 电流采用专用短接片、CT侧短接后打开中连片并红胶布隔离，符合5.2.3a2；220kV/110kV电流A/B/C/N四相齐全。")
P("✓ 电压回路整体覆盖220kV/110kV I、II母及35kV，A/B/C/L四相，A/B屏分开，规模完整。")
H("（六）信号回路（第八章 / 5.2.4）",3)
P("■ 问题：缺测控→保护屏信号回路公共端解除（S4）；遥信电源空开未标封挡（M6）。")
P("✓ 故障录波器至主变保护A/B屏录波公共端（G900/2A、G900/2B）已拆除并红胶布包好。")
H("（七）交直流电源（第九章 / 5.2.5）",3)
P("■ 问题：直流/遥信空开未标封挡（M6）；缺交流电源回路（M7）。")
P("✓ 直流空开（保护A柜/非电量/高压侧操作箱2，I/II分屏）已列明并断开。")
H("（八）过程安全措施（第十章 / 5.2.6）",3)
P("■ 问题：缺过程安措章节（M8）。若含接入母差传动试验，须补充母差所有硬压板退出并封挡。")

# ===== 四、结论与签字 =====
doc.add_paragraph()
H("四、审查结论",1)
P("结论：修改后复审。本票电流、电压主体隔离措施框架基本完整，但存在 1 项高危电压端子复制错误（M6，可能误拆220kV运行电压）、1 项电流中性线标错（M5），以及封挡标注、交流电源、遥控电源、过程安措、票面要素等多项缺失，须整改并复核后方可签发、执行。", bold=True, color=RGBColor(0xC0,0x00,0x00))
P("")
ct=doc.add_table(rows=4,cols=4); ct.style='Table Grid'
for i,row in enumerate([("审查岗位","审查人","审查日期","结论"),("编制自审","","",""),("技术负责人审核","","",""),("签发人签发","","","")]):
    c=ct.rows[i].cells
    if i==0:
        for j,txt in enumerate(row): cell(c[j],txt,True,10,fill="D9E2F3")
    else:
        for j,txt in enumerate(row): cell(c[j],txt,False,10)

doc.save(OUT)
print("Saved:", OUT)
