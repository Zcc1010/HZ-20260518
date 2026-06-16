import re
import subprocess
from pathlib import Path

import fitz
import xlrd


def convert_to_md(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = path.suffix.lower()

    if suffix in (".md", ".txt"):
        return path.read_text(encoding="utf-8")
    elif suffix == ".xls":
        return _xls_to_md(path)
    elif suffix == ".xlsx":
        return _xlsx_to_md(path)
    elif suffix == ".doc":
        return _doc_to_md(path)
    elif suffix == ".docx":
        return _docx_to_md(path)
    elif suffix == ".pdf":
        return _pdf_to_md(path)
    else:
        raise ValueError(f"不支持的文件格式: {suffix}")


def _clean_cell(val) -> str:
    if isinstance(val, float) and val == int(val):
        text = str(int(val))
    else:
        text = str(val).strip()
    text = text.replace("\n", "").replace("\r", "")
    return text


def _xls_to_md(path: Path) -> str:
    wb = xlrd.open_workbook(str(path), formatting_info=True)
    parts = []
    for si in range(len(wb.sheet_names())):
        ws = wb.sheet_by_index(si)
        parts.append(f"## Sheet: {wb.sheet_names()[si]}\n")

        raw_grid = []
        for r in range(ws.nrows):
            row = [_clean_cell(ws.cell_value(r, c)) for c in range(ws.ncols)]
            raw_grid.append(row)

        for crlo, crhi, cclo, cchi in ws.merged_cells:
            for r in range(crlo, crhi):
                for c in range(cclo, cchi):
                    if r != crlo or c != cclo:
                        raw_grid[r][c] = ""

        max_col = 0
        for row in raw_grid:
            for c in range(len(row) - 1, -1, -1):
                if row[c]:
                    max_col = max(max_col, c + 1)
                    break

        for row in raw_grid:
            if any(row[:max_col]):
                parts.append(" | ".join(row[:max_col]))
        parts.append("")
    return "\n".join(parts)


def _xlsx_to_md(path: Path) -> str:
    import datetime
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True)
    parts = []
    for name in wb.sheetnames:
        ws = wb[name]
        parts.append(f"## Sheet: {name}\n")

        raw_grid = []
        for row in ws.iter_rows(min_row=1, max_row=ws.max_row, max_col=ws.max_column, values_only=True):
            cells = []
            for c in row:
                if c is None:
                    cells.append("")
                elif isinstance(c, datetime.datetime):
                    cells.append(c.strftime("%Y-%m-%d") if c.hour == 0 and c.minute == 0 and c.second == 0 else str(c))
                elif isinstance(c, float) and c == int(c):
                    cells.append(str(int(c)))
                else:
                    cells.append(str(c).strip())
            raw_grid.append(cells)

        max_col = 0
        for row in raw_grid:
            for c in range(len(row) - 1, -1, -1):
                if row[c]:
                    max_col = max(max_col, c + 1)
                    break

        for row in raw_grid:
            if any(row[:max_col]):
                parts.append(" | ".join(row[:max_col]))
        parts.append("")
    return "\n".join(parts)


def _doc_to_md(path: Path) -> str:
    # Try antiword first (fast, accurate for MS Word .doc)
    try:
        result = subprocess.run(
            ["antiword", str(path)],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            lines = result.stdout.splitlines()
            if lines and lines[0].startswith("convert "):
                lines = lines[1:]
            return "\n".join(lines)
    except FileNotFoundError:
        pass  # antiword not installed, fall through

    # Fallback: pure Python OLE2 text extraction (works with WPS .doc)
    return _doc_to_md_olefile(path)


def _doc_to_md_olefile(path: Path) -> str:
    import struct

    try:
        import olefile
    except ImportError:
        raise RuntimeError(
            "antiword 不可用且未安装 olefile。"
            "请安装 antiword 或运行: pip install olefile"
        )

    ole = olefile.OleFileIO(str(path))
    wd = ole.openstream("WordDocument").read()

    wIdent = struct.unpack_from("<H", wd, 0)[0]
    if wIdent != 0xA5EC:
        ole.close()
        raise ValueError(f"不是有效的 Word 文档: {path.name}")

    # Scan WordDocument stream for UTF-16LE encoded text
    # WPS .doc files store text as UTF-16LE in the main stream
    results = []
    i = 0x200  # skip FIB header
    current: list[str] = []
    while i < len(wd) - 1:
        char = struct.unpack_from("<H", wd, i)[0]
        if (
            (0x20 <= char <= 0x7E)
            or (0x4E00 <= char <= 0x9FFF)  # CJK Unified
            or (0x3000 <= char <= 0x303F)  # CJK Symbols
            or (0xFF00 <= char <= 0xFFEF)  # Fullwidth Forms
            or char in (0x000A, 0x000D, 0x0009)
        ):
            current.append(chr(char))
        else:
            if len(current) >= 5:
                results.append("".join(current))
            current = []
        i += 2

    if len(current) >= 5:
        results.append("".join(current))

    ole.close()

    text = "\n".join(results)
    # Remove control characters except newline/tab
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
    return text


def _docx_to_md(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
        parts.append("")
    return "\n".join(parts)


def _pdf_to_md(path: Path) -> str:
    doc = fitz.open(str(path))
    parts = []
    for page in doc:
        tabs = page.find_tables()
        table_bboxes = [t.bbox for t in tabs.tables]

        blocks = page.get_text("blocks")
        elements = []
        for b in blocks:
            x0, y0, x1, y1, text, *_ = b
            text = text.strip().replace("\n", " ")
            if not text:
                continue
            inside_table = any(
                bx0 <= x0 and y0 >= by0 and y1 <= by1
                for bx0, by0, bx1, by1 in table_bboxes
            )
            if not inside_table:
                elements.append((y0, "text", text))

        for t in tabs.tables:
            ty0 = t.bbox[1]
            rows = []
            for row in t.extract():
                cells = [str(c or "").replace("\n", " ").strip() for c in row]
                if any(cells):
                    rows.append(" | ".join(cells))
            if rows:
                elements.append((ty0, "table", rows))

        elements.sort(key=lambda e: e[0])

        for _, typ, content in elements:
            if typ == "text":
                parts.append(content)
            else:
                parts.extend(content)
            parts.append("")
    return "\n".join(parts)
