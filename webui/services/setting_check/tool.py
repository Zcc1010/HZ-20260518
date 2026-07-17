"""Setting check read/write tools — read and modify setting check reports."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from loguru import logger

from nanobot.agent.tools.base import Tool, tool_parameters
from nanobot.agent.tools.schema import StringSchema, tool_parameters_schema


_AGENTPLAYGROUND_DIR = (Path.home() / ".nanobot" / "agentplayground").resolve()


def _check_path(path: Path) -> None:
    """确保文件路径在 agentplayground 目录下，防止越权写入。"""
    resolved = path.resolve()
    try:
        resolved.relative_to(_AGENTPLAYGROUND_DIR)
    except ValueError:
        raise PermissionError(f"路径 {path} 不在 agentplayground 目录下，拒绝操作")


def _get_service():
    from webui.services.setting_check.service import SettingCheckService
    return SettingCheckService()


def _find_report_file(job_root: Path) -> Path | None:
    """在 job output 目录中查找报告 .md 文件。"""
    output_dir = job_root / "output"
    if not output_dir.exists():
        return None
    # 优先匹配 定值校核报告.md
    preferred = list(output_dir.rglob("*定值校核报告.md"))
    if preferred:
        return preferred[0]
    # 回退：任意 .md
    md_files = list(output_dir.rglob("*.md"))
    return md_files[0] if md_files else None


def _replace_section(full_content: str, section_keyword: str, new_section_body: str) -> str | None:
    """替换 Markdown 文档中指定章节的内容。

    Args:
        full_content: 完整的 Markdown 文档
        section_keyword: 章节标题关键字（部分匹配）
        new_section_body: 该章节的新内容（不含 ## 标题行）

    Returns:
        替换后的完整内容，未找到章节时返回 None
    """
    import re

    # 按 ## 标题分割，保留标题行
    parts = re.split(r"(?=^## )", full_content, flags=re.MULTILINE)

    for i, part in enumerate(parts):
        lines = part.strip().split("\n", 1)
        if not lines:
            continue
        title_line = lines[0]
        if not title_line.startswith("## "):
            continue
        if section_keyword in title_line:
            # 找到目标章节，替换内容
            new_part = title_line + "\n\n" + new_section_body.strip() + "\n"
            parts[i] = new_part
            return "\n".join(parts)

    return None


READ_TOOL_DESC = """读取定值校核报告内容。可通过 job_id 直接读取，或通过站点名称搜索匹配的任务。

参数说明：
- job_id: 任务 ID，直接指定要读取的任务
- station: 站点名称，模糊搜索该站点下的定值校核任务

至少需要提供 job_id 或 station 其中一个参数。
"""

WRITE_TOOL_DESC = """改写定值校核报告的指定章节。只替换目标章节内容，其余章节保持不变。

参数说明：
- job_id: 任务 ID（必填）
- section: 要修改的章节标题（必填），例如"基本信息"、"校核结论"、"存在问题"等，支持部分匹配
- content: 该章节的新内容（必填，Markdown 格式，不含 ## 标题行）

使用流程：先用 setting_check_read 读取报告，找到要修改的章节标题，再用本工具只替换该章节。
"""


@tool_parameters(
    tool_parameters_schema(
        job_id=StringSchema("任务 ID，直接指定要读取的任务"),
        station=StringSchema("站点名称，模糊搜索该站点下的定值校核任务"),
    )
)
class SettingCheckReadTool(Tool):
    """读取定值校核报告内容。支持按 job_id 或站点名称查询。"""

    @property
    def name(self) -> str:
        return "setting_check_read"

    @property
    def description(self) -> str:
        return READ_TOOL_DESC

    @property
    def read_only(self) -> bool:
        return True

    async def execute(self, **kwargs: Any) -> str:
        job_id = (kwargs.get("job_id") or "").strip()
        station = (kwargs.get("station") or "").strip()

        if not job_id and not station:
            return "错误：请至少提供一个参数：job_id 或 station"

        service = _get_service()

        # 按 job_id 直接读取
        if job_id:
            job = service.get_job(job_id)
            if not job:
                return f"错误：未找到任务 ID 为 '{job_id}' 的定值校核任务"
            return self._read_report(service, job)

        # 按 station 搜索（SettingCheckService 无 search_jobs，手动筛选）
        all_jobs = service.list_jobs()
        matched = [j for j in all_jobs if station in (j.get("station") or "")]
        if not matched:
            return f"错误：未找到站点 '{station}' 相关的定值校核任务"

        completed = [j for j in matched if j.get("status") == "completed"]
        if not completed:
            names = [f"  - {j.get('station', '')} / {j.get('device', '')}（{j['id']}）状态: {j.get('status', '')}" for j in matched]
            return f"找到 {len(matched)} 个任务，但均未完成：\n" + "\n".join(names)

        # 如果只有一个已完成任务，直接返回报告
        if len(completed) == 1:
            return self._read_report(service, completed[0])

        # 多个任务，返回列表供选择
        lines = [f"找到 {len(completed)} 个已完成的定值校核任务："]
        for j in completed:
            lines.append(f"  - ID: {j['id']}  站点: {j.get('station', '')}  设备: {j.get('device', '')}  时间: {j.get('created_at', '')}")
        lines.append("\n请使用 job_id 参数指定要读取的任务。")
        return "\n".join(lines)

    def _read_report(self, service, job: dict) -> str:
        job_id = job["id"]
        report_path = _find_report_file(service.app_root / "jobs" / job_id)
        if not report_path:
            return f"错误：任务 {job_id} 已完成但未找到报告文件"

        _check_path(report_path)

        try:
            content = report_path.read_text(encoding="utf-8")
        except Exception as exc:
            return f"错误：读取报告文件失败：{exc}"

        station = job.get("station", "未知")
        device = job.get("device", "未知")
        header = f"定值校核报告 | {station} / {device}\n任务 ID: {job_id}\n"
        header += "=" * 60 + "\n\n"
        footer = f"\n\n---\n如需修改此报告，请使用 setting_check_write 工具。只需指定要修改的章节，其余章节自动保留。示例：setting_check_write(job_id=\"{job_id}\", section=\"校核结论\", content=\"该章节的新内容\")。不要使用 edit_file 或 write_file。"
        return header + content + footer


@tool_parameters(
    tool_parameters_schema(
        job_id=StringSchema("任务 ID（必填）"),
        section=StringSchema("要修改的章节标题关键字（必填），例如：基本信息、校核结论、存在问题"),
        content=StringSchema("该章节的新内容（必填），Markdown 格式，不含 ## 标题行"),
    )
)
class SettingCheckWriteTool(Tool):
    """改写定值校核报告的指定章节。只替换目标章节，其余章节保持不变。"""

    @property
    def name(self) -> str:
        return "setting_check_write"

    @property
    def description(self) -> str:
        return WRITE_TOOL_DESC

    @property
    def read_only(self) -> bool:
        return False

    async def execute(self, **kwargs: Any) -> str:
        job_id = (kwargs.get("job_id") or "").strip()
        section = (kwargs.get("section") or "").strip()
        content = kwargs.get("content") or ""

        if not job_id:
            return "错误：请提供 job_id 参数"
        if not section:
            return "错误：请提供 section 参数（要修改的章节标题关键字）"
        if content is None:
            return "错误：请提供 content 参数（该章节的新内容，允许为空字符串）"

        service = _get_service()
        job = service.get_job(job_id)
        if not job:
            return f"错误：未找到任务 ID 为 '{job_id}' 的定值校核任务"

        report_path = _find_report_file(service.app_root / "jobs" / job_id)
        if not report_path:
            return f"错误：任务 {job_id} 未找到报告文件，无法写入"

        _check_path(report_path)

        try:
            existing = report_path.read_text(encoding="utf-8")
        except Exception as exc:
            return f"错误：读取报告文件失败：{exc}"

        new_full = _replace_section(existing, section, content)
        if new_full is None:
            import re
            headers = re.findall(r"^## .+", existing, re.MULTILINE)
            header_list = "\n".join(f"  {h}" for h in headers) if headers else "  （无）"
            return f"错误：未找到包含 '{section}' 的章节。可用章节：\n{header_list}"

        try:
            report_path.write_text(new_full, encoding="utf-8")
        except Exception as exc:
            return f"错误：写入报告文件失败：{exc}"

        # 同步更新 .docx
        docx_path = report_path.with_suffix(".docx")
        try:
            from webui.utils.md_to_docx import MdToDocxConverter
            converter = MdToDocxConverter()
            converter.convert(new_full, docx_path)
        except Exception:
            pass

        station = job.get("station", "未知")
        device = job.get("device", "未知")
        return f"定值校核报告「{section}」章节已更新：{station} / {device}（{job_id}）"


RERUN_TOOL_DESC = """使用原始输入文件重新执行定值校核。会覆盖原有报告。

参数说明：
- job_id: 任务 ID（必填）

重新校核会读取任务的原始定值单和计算书文件，重新运行校核 pipeline 生成新报告。
"""

GENERATE_TOOL_DESC = """根据工作区文件直接生成定值校核报告。

参数说明：
- workspace: 工作区名称（必填），如 "安徽.阳湖变-测试"

工具会自动：
1. 从工作区的 定值单/ 目录读取所有定值单文件
2. 从工作区的 计算书/ 目录读取计算书文件
3. 执行定值校核 pipeline 生成报告
4. 将报告保存到工作区的 报告/ 目录

使用场景：用户要求"重新生成报告"、"生成校核报告"时使用此工具。
"""


@tool_parameters(
    tool_parameters_schema(
        job_id=StringSchema("任务 ID（必填）"),
    )
)
class SettingCheckRerunTool(Tool):
    """使用原始输入文件重新执行定值校核。"""

    @property
    def name(self) -> str:
        return "setting_check_rerun"

    @property
    def description(self) -> str:
        return RERUN_TOOL_DESC

    @property
    def read_only(self) -> bool:
        return False

    async def execute(self, **kwargs: Any) -> str:
        import asyncio
        import json

        job_id = (kwargs.get("job_id") or "").strip()
        if not job_id:
            return "错误：请提供 job_id 参数"

        service = _get_service()
        job = service.get_job(job_id)
        if not job:
            return f"错误：未找到任务 ID 为 '{job_id}' 的定值校核任务"

        job_root = service.app_root / "jobs" / job_id
        _check_path(job_root)

        manifest_path = job_root / "inputs.json"
        if not manifest_path.exists():
            return f"错误：任务 {job_id} 缺少 inputs.json，无法重新校核"

        try:
            manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        except Exception as exc:
            return f"错误：读取 inputs.json 失败：{exc}"

        inputs_dir = job_root / "inputs"
        setting_names = manifest.get("setting_files", [])
        calc_names = manifest.get("calc_files", manifest.get("calc_file", []))

        if isinstance(calc_names, str):
            calc_names = [calc_names]

        setting_paths = [inputs_dir / n for n in setting_names]
        calc_paths = [inputs_dir / n for n in calc_names]

        missing = [str(p) for p in setting_paths + calc_paths if not p.exists()]
        if missing:
            return f"错误：原始文件缺失，无法重新校核：\n" + "\n".join(f"  {m}" for m in missing)

        # 标记任务为处理中，前端可通过轮询获取进度
        service.mark_processing(job_id)
        service.update_progress(job_id, 0, "开始重新校核...")

        def _progress(progress: int, message: str) -> None:
            service.update_progress(job_id, progress, message)

        try:
            from webui.services.setting_check.service import execute_setting_check
            await asyncio.to_thread(execute_setting_check, job_root, setting_paths, calc_paths, _progress)
        except Exception as exc:
            service.update_progress(job_id, 0, f"失败：{exc}")
            return f"错误：重新校核失败：{exc}"

        # 同步更新 .docx
        new_report = _find_report_file(job_root)
        docx_path = None
        if new_report:
            docx_path = new_report.with_suffix(".docx")
            try:
                from webui.utils.md_to_docx import MdToDocxConverter
                converter = MdToDocxConverter()
                converter.convert(new_report.read_text(encoding="utf-8"), docx_path)
            except Exception:
                pass

        # 标记任务完成，前端轮询会更新状态
        result_file = docx_path if docx_path and docx_path.exists() else new_report
        if result_file:
            try:
                service.mark_completed(job_id, result_file)
            except Exception:
                pass

        station = job.get("station", "未知")
        device = job.get("device", "未知")
        result = f"定值校核已重新执行：{station} / {device}（{job_id}）"
        files = []
        if docx_path and docx_path.exists():
            files.append(str(docx_path))
        elif new_report:
            files.append(str(new_report))
        if files:
            result += f"\n【必须调用 message 工具发送文件】调用方式：message(content=\"定值校核报告已重新生成\", media={files})。不要把文件路径作为链接展示，必须通过 message 工具发送。"
        return result


@tool_parameters(
    tool_parameters_schema(
        workspace=StringSchema("工作区名称（必填），如 '安徽.阳湖变-测试'"),
    )
)
class SettingCheckGenerateTool(Tool):
    """根据工作区文件直接生成定值校核报告。"""

    @property
    def name(self) -> str:
        return "setting_check_generate"

    @property
    def description(self) -> str:
        return GENERATE_TOOL_DESC

    @property
    def read_only(self) -> bool:
        return False

    async def execute(self, **kwargs: Any) -> str:
        import asyncio
        import re

        workspace = (kwargs.get("workspace") or "").strip()
        if not workspace:
            return "错误：请提供 workspace 参数（工作区名称）"

        # 如果传入的是完整路径，提取最后的工作区名称
        if "/" in workspace or "\\" in workspace:
            workspace = workspace.rstrip("/").rstrip("\\").split("/")[-1].split("\\")[-1]

        # 去除名称中多余的空格（如 "施官变 - 独施 369" → "施官变-独施369"）
        workspace = re.sub(r'\s*-\s*', '-', workspace)
        workspace = re.sub(r'\s+', '', workspace)

        # 工作区路径
        ws_path = _AGENTPLAYGROUND_DIR / "setting-check" / "workspace" / workspace
        if not ws_path.exists():
            # 尝试模糊匹配
            parent = _AGENTPLAYGROUND_DIR / "setting-check" / "workspace"
            if parent.exists():
                candidates = [d.name for d in parent.iterdir() if d.is_dir()]
                # 去除空格后匹配
                normalized = workspace.replace(" ", "")
                for c in candidates:
                    if c.replace(" ", "") == normalized:
                        ws_path = parent / c
                        workspace = c
                        break
            if not ws_path.exists():
                return f"错误：工作区 '{workspace}' 不存在。可用工作区：{', '.join(candidates) if 'candidates' in dir() else '无'}"

        # 检查定值单目录
        setting_dir = ws_path / "定值单"
        if not setting_dir.exists() or not any(setting_dir.iterdir()):
            return f"错误：工作区 '{workspace}' 的 定值单/ 目录不存在或为空"

        # 检查计算书目录
        calc_dir = ws_path / "计算书"
        if not calc_dir.exists() or not any(calc_dir.iterdir()):
            return f"错误：工作区 '{workspace}' 的 计算书/ 目录不存在或为空"

        # 收集定值单文件
        setting_paths = []
        for f in sorted(setting_dir.iterdir()):
            if f.is_file() and f.suffix.lower() in {'.xls', '.xlsx', '.doc', '.docx', '.pdf', '.md', '.txt'}:
                setting_paths.append(f)

        if not setting_paths:
            return f"错误：工作区 '{workspace}' 的 定值单/ 目录中未找到有效文件"

        # 收集计算书文件
        calc_paths = []
        for f in sorted(calc_dir.iterdir()):
            if f.is_file() and f.suffix.lower() in {'.xls', '.xlsx', '.doc', '.docx', '.pdf', '.md', '.txt'}:
                calc_paths.append(f)

        if not calc_paths:
            return f"错误：工作区 '{workspace}' 的 计算书/ 目录中未找到有效文件"

        # 创建临时输出目录
        import tempfile
        temp_dir = Path(tempfile.mkdtemp())
        output_dir = temp_dir / "output"
        output_dir.mkdir(parents=True, exist_ok=True)

        try:
            # 执行校核 pipeline
            from webui.services.setting_check.pipeline import run_pipeline
            from webui.trip_briefing.llm.client import LLMClient
            import json

            # 获取 LLM 配置
            config_path = Path.home() / ".protection" / "config.json"
            if not config_path.exists():
                config_path = Path.home() / ".nanobot" / "config.json"

            if not config_path.exists():
                return "错误：缺少配置文件 config.json，无法调用 LLM"

            config_data = json.loads(config_path.read_text(encoding="utf-8"))
            default_model = config_data.get("agents", {}).get("defaults", {}).get("model", "glm-4-flash")

            providers = config_data.get("providers", {})
            provider = None

            # 各 provider 的合理默认模型
            provider_default_models = {
                "zhipu": "glm-4-flash",
                "dashscope": "qwen-turbo",
                "deepseek": "deepseek-chat",
                "openai": "gpt-4o-mini",
                "openrouter": "openai/gpt-4o-mini",
                "custom": default_model,
            }

            preferred_providers = ["custom", "zhipu", "dashscope", "deepseek", "openai", "openrouter"]
            for name in preferred_providers:
                p = providers.get(name, {})
                if p.get("apiKey") or p.get("api_key"):
                    fallback_model = provider_default_models.get(name, default_model)
                    provider = {
                        "base_url": p.get("apiBase") or p.get("base_url", ""),
                        "api_key": p.get("apiKey") or p.get("api_key", ""),
                        "model": p.get("model", fallback_model),
                    }
                    break

            if not provider:
                for name, p in providers.items():
                    if p.get("apiKey") or p.get("api_key"):
                        fallback_model = provider_default_models.get(name, default_model)
                        provider = {
                            "base_url": p.get("apiBase") or p.get("base_url", ""),
                            "api_key": p.get("apiKey") or p.get("api_key", ""),
                            "model": p.get("model", fallback_model),
                        }
                        break

            if not provider:
                return "错误：配置文件中未找到有效的 provider"

            api_url = provider.get("base_url", "")
            api_key = provider.get("api_key", "")
            model = provider.get("model", default_model)

            if not api_key:
                return "错误：配置文件中未找到 API key"

            llm_client = LLMClient(
                api_url=api_url,
                api_key=api_key,
                model=model,
                timeout=180,
                max_retries=3,
                enable_thinking=False,
            )

            def llm_call(prompt: str) -> str:
                response = llm_client.chat_completion(
                    messages=[{"role": "user", "content": prompt}],
                    model=model,
                    max_tokens=8192,
                )
                if not response.success:
                    raise RuntimeError(f"LLM 调用失败: {response.error_message}")
                return response.content

            # 运行 pipeline
            logger.info("[setting_check_generate] 开始运行校核 pipeline, workspace={}", workspace)
            result = run_pipeline(
                setting_paths=[str(p) for p in setting_paths],
                calc_paths=[str(p) for p in calc_paths],
                llm_call_func=llm_call,
                output_dir=str(output_dir),
            )

            report_path = Path(result["report_path"])

            # 复制报告到工作区的报告目录
            report_dir = ws_path / "报告"
            report_dir.mkdir(parents=True, exist_ok=True)

            # 目标文件名
            target_name = f"{workspace}定值校核报告.md"
            target_path = report_dir / target_name

            # 复制并覆盖
            import shutil
            shutil.copy2(str(report_path), str(target_path))

            # 同步生成 .docx
            docx_path = target_path.with_suffix(".docx")
            try:
                from webui.utils.md_to_docx import MdToDocxConverter
                converter = MdToDocxConverter()
                converter.convert(target_path.read_text(encoding="utf-8"), docx_path)
            except Exception:
                pass

            # 清理临时目录
            shutil.rmtree(temp_dir, ignore_errors=True)

            station = result.get("device_info", {}).get("station", workspace)
            device = result.get("device_info", {}).get("device", "")

            return f"定值校核报告已生成，可在左侧「报告」目录下查看。"

        except Exception as exc:
            # 清理临时目录
            import shutil
            shutil.rmtree(temp_dir, ignore_errors=True)
            return f"错误：生成报告失败：{exc}"


WORKSPACE_READ_TOOL_DESC = """读取定值校核工作区中的文件内容。可读取定值单、计算书、说明书、整定原则、台账、报告等目录下的文件。

参数说明：
- workspace: 工作区名称（必填），如 "安徽.阳湖变-测试"
- path: 文件相对路径（必填），如 "定值单/xxx.xlsx"、"计算书/xxx.docx"、"说明书/xxx.pdf"

支持的文件类型：
- 文本文件（.md, .txt）：直接返回文本内容
- Excel 文件（.xlsx, .xls）：转换为 CSV 文本返回
- Word 文件（.docx）：提取文本内容返回
- PDF 文件（.pdf）：提取文本内容返回
- 图片文件（.png, .jpg 等）：返回说明信息

使用场景：用户要求重新审核、查看工作区文件、基于工作区文件分析时，先用此工具读取相关文件内容。
"""


@tool_parameters(
    tool_parameters_schema(
        workspace=StringSchema("工作区名称（必填），如 '安徽.阳湖变-测试'"),
        path=StringSchema("文件相对路径（必填），如 '定值单/xxx.xlsx'、'计算书/xxx.docx'"),
    )
)
class SettingCheckWorkspaceReadTool(Tool):
    """读取定值校核工作区中的文件内容。"""

    @property
    def name(self) -> str:
        return "setting_check_workspace_read"

    @property
    def description(self) -> str:
        return WORKSPACE_READ_TOOL_DESC

    @property
    def read_only(self) -> bool:
        return True

    async def execute(self, **kwargs: Any) -> str:
        import re

        workspace = (kwargs.get("workspace") or "").strip()
        path = (kwargs.get("path") or "").strip()

        if not workspace:
            return "错误：请提供 workspace 参数（工作区名称）"
        if not path:
            return "错误：请提供 path 参数（文件相对路径，如 '定值单/xxx.xlsx'）"

        # 标准化工作区名称
        if "/" in workspace or "\\" in workspace:
            workspace = workspace.rstrip("/").rstrip("\\").split("/")[-1].split("\\")[-1]
        workspace = re.sub(r'\s*-\s*', '-', workspace)
        workspace = re.sub(r'\s+', '', workspace)

        ws_path = _AGENTPLAYGROUND_DIR / "setting-check" / "workspace" / workspace
        if not ws_path.exists():
            # 模糊匹配
            parent = _AGENTPLAYGROUND_DIR / "setting-check" / "workspace"
            if parent.exists():
                candidates = [d.name for d in parent.iterdir() if d.is_dir()]
                normalized = workspace.replace(" ", "")
                for c in candidates:
                    if c.replace(" ", "") == normalized:
                        ws_path = parent / c
                        workspace = c
                        break
            if not ws_path.exists():
                return f"错误：工作区 '{workspace}' 不存在"

        # 安全拼接路径
        try:
            full = _safe_join(ws_path, path)
        except ValueError:
            return f"错误：路径不合法或越权访问"

        if not full.exists():
            return f"错误：文件不存在：{path}"

        if full.is_dir():
            # 返回目录下的文件列表
            items = []
            for entry in sorted(full.iterdir(), key=lambda e: (not e.is_dir(), e.name.lower())):
                if entry.name.startswith("."):
                    continue
                kind = "[目录]" if entry.is_dir() else "[文件]"
                items.append(f"  {kind} {entry.name}")
            return f"「{path}」是一个目录，包含 {len(items)} 个项目：\n" + "\n".join(items)

        ext = full.suffix.lower()
        size = full.stat().st_size

        # 限制文件大小（5MB）
        if size > 5 * 1024 * 1024:
            return f"错误：文件过大（{size / 1024 / 1024:.1f}MB），超过 5MB 限制"

        # 图片：返回说明
        if ext in {'.png', '.jpg', '.jpeg', '.gif', '.svg', '.webp'}:
            return f"「{path}」是一个图片文件（{ext}，{size / 1024:.1f}KB）。如需查看图片内容，请使用 read_file 工具读取。"

        # 文本文件：直接读取
        if ext in {'.md', '.txt', '.csv', '.json', '.xml', '.html', '.htm'}:
            try:
                text = full.read_text(encoding="utf-8")
            except UnicodeDecodeError:
                text = full.read_text(encoding="gbk", errors="replace")
            return f"文件：{workspace}/{path}\n大小：{size / 1024:.1f}KB\n{'=' * 60}\n\n{text}"

        # Excel 文件：转换为 CSV
        if ext in {'.xlsx', '.xls'}:
            try:
                import openpyxl
                wb = openpyxl.load_workbook(str(full), read_only=True, data_only=True)
                output_parts = []
                for sheet_name in wb.sheetnames:
                    ws_obj = wb[sheet_name]
                    rows = []
                    for row in ws_obj.iter_rows(values_only=True):
                        cells = [str(c) if c is not None else "" for c in row]
                        rows.append(",".join(cells))
                    sheet_text = "\n".join(rows)
                    output_parts.append(f"### Sheet: {sheet_name}\n{sheet_text}")
                wb.close()
                content = "\n\n".join(output_parts)
                return f"文件：{workspace}/{path}\n类型：Excel\n工作表：{', '.join(wb.sheetnames)}\n{'=' * 60}\n\n{content}"
            except Exception as exc:
                return f"错误：读取 Excel 文件失败：{exc}"

        # Word 文件：提取文本
        if ext in {'.docx', '.doc'}:
            try:
                if ext == '.docx':
                    from docx import Document
                    doc = Document(str(full))
                    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                    content = "\n\n".join(paragraphs)
                    # 也提取表格内容
                    for table in doc.tables:
                        for row in table.rows:
                            cells = [cell.text.strip() for cell in row.cells]
                            content += "\n" + " | ".join(cells)
                else:
                    # .doc 格式尝试 antiword
                    import subprocess
                    result = subprocess.run(
                        ["antiword", str(full)],
                        capture_output=True, text=True, timeout=30,
                    )
                    if result.returncode == 0:
                        content = result.stdout
                    else:
                        return f"错误：无法读取 .doc 格式文件，请转换为 .docx 后重试"
                return f"文件：{workspace}/{path}\n类型：Word 文档\n{'=' * 60}\n\n{content}"
            except Exception as exc:
                return f"错误：读取 Word 文件失败：{exc}"

        # PDF 文件：提取文本
        if ext == '.pdf':
            try:
                import subprocess
                # 尝试 pdftotext
                result = subprocess.run(
                    ["pdftotext", str(full), "-"],
                    capture_output=True, text=True, timeout=60,
                )
                if result.returncode == 0 and result.stdout.strip():
                    content = result.stdout
                    return f"文件：{workspace}/{path}\n类型：PDF\n{'=' * 60}\n\n{content}"
                else:
                    return f"错误：无法提取 PDF 文本内容（pdftotext 不可用或 PDF 为扫描件）"
            except FileNotFoundError:
                return f"错误：系统未安装 pdftotext 工具，无法读取 PDF 文件"
            except Exception as exc:
                return f"错误：读取 PDF 文件失败：{exc}"

        # 其他文件：尝试读取为文本
        try:
            text = full.read_text(encoding="utf-8", errors="replace")
            return f"文件：{workspace}/{path}\n类型：{ext}\n{'=' * 60}\n\n{text}"
        except Exception:
            return f"错误：不支持的文件类型：{ext}"


def _safe_join(base: Path, target: str) -> Path:
    """防止路径穿越，允许软链接指向 resources 和 temp 目录。"""
    result = (base / target).resolve()
    base_resolved = base.resolve()
    if str(result).startswith(str(base_resolved)):
        return result
    resources_dir = (_AGENTPLAYGROUND_DIR / "resources").resolve()
    if str(result).startswith(str(resources_dir)):
        return result
    temp_dir = (_AGENTPLAYGROUND_DIR / "temp").resolve()
    if str(result).startswith(str(temp_dir)):
        return result
    raise ValueError("path traversal detected")
