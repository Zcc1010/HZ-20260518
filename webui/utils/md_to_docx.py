# -*- coding: utf-8 -*-
"""Markdown 转 Word (.docx) 转换工具"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# LaTeX -> Unicode 映射
LATEX_UNICODE = {
    r"\times": "×",
    r"\geq": "≥",
    r"\leq": "≤",
    r"\neq": "≠",
    r"\approx": "≈",
    r"\pm": "±",
    r"\mp": "∓",
    r"\cdot": "·",
    r"\ldots": "…",
    r"\cdots": "⋯",
    r"\alpha": "α",
    r"\beta": "β",
    r"\gamma": "γ",
    r"\delta": "δ",
    r"\epsilon": "ε",
    r"\zeta": "ζ",
    r"\eta": "η",
    r"\theta": "θ",
    r"\iota": "ι",
    r"\kappa": "κ",
    r"\lambda": "λ",
    r"\mu": "μ",
    r"\nu": "ν",
    r"\xi": "ξ",
    r"\pi": "π",
    r"\rho": "ρ",
    r"\sigma": "σ",
    r"\tau": "τ",
    r"\upsilon": "υ",
    r"\phi": "φ",
    r"\chi": "χ",
    r"\psi": "ψ",
    r"\omega": "ω",
    r"\Alpha": "Α",
    r"\Beta": "Β",
    r"\Gamma": "Γ",
    r"\Delta": "Δ",
    r"\Theta": "Θ",
    r"\Lambda": "Λ",
    r"\Xi": "Ξ",
    r"\Pi": "Π",
    r"\Sigma": "Σ",
    r"\Phi": "Φ",
    r"\Psi": "Ψ",
    r"\Omega": "Ω",
    r"\infty": "∞",
    r"\partial": "∂",
    r"\nabla": "∇",
    r"\sum": "∑",
    r"\prod": "∏",
    r"\int": "∫",
    r"\rightarrow": "→",
    r"\leftarrow": "←",
    r"\Rightarrow": "⇒",
    r"\Leftarrow": "⇐",
    r"\leftrightarrow": "↔",
    r"\sim": "∼",
    r"\propto": "∝",
    r"\perp": "⊥",
    r"\parallel": "∥",
    r"\angle": "∠",
    r"\degree": "°",
}

GREEK_LOWER = set("αβγδεζηθικλμνξπρστυφχψω")
GREEK_UPPER = set("ΑΒΓΔΘΛΞΠΣΦΨΩ")


def _convert_latex(text: str) -> str:
    """将 LaTeX 数学表达式转换为可读的 Unicode 文本。"""
    # 替换已知命令
    for cmd, char in LATEX_UNICODE.items():
        text = text.replace(cmd, char)
    # \frac{a}{b} -> a/b
    text = re.sub(r"\\frac\{([^}]+)\}\{([^}]+)\}", r"\1/\2", text)
    # \sqrt{x} -> √x
    text = re.sub(r"\\sqrt\{([^}]+)\}", r"√\1", text)
    # \, -> 空格
    text = text.replace(r"\,", " ")
    text = text.replace(r"\;", " ")
    text = text.replace(r"\quad", "  ")
    text = text.replace(r"\qquad", "    ")
    # _{text} -> _text
    text = re.sub(r"_\{([^}]+)\}", r"_\1", text)
    # ^{text} -> ^text
    text = re.sub(r"\^\{([^}]+)\}", r"^\1", text)
    # 移除剩余的 \text{} 包装
    text = re.sub(r"\\text\{([^}]+)\}", r"\1", text)
    # 移除剩余的反斜杠命令（未知命令）
    text = re.sub(r"\\[a-zA-Z]+", "", text)
    # 清理多余空格
    text = re.sub(r"  +", " ", text)
    return text.strip()


def _parse_inline(paragraph, text: str, font_name: str = "Microsoft YaHei",
                   font_size: Pt = Pt(11)):
    """解析行内 Markdown（加粗、公式等）并添加到段落。"""
    # 分割：加粗 **...** 和行内公式 $...$
    parts = re.split(r"(\*\*.*?\*\*|\$\$.*?\$\$|\$[^$]+\$)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            # 加粗
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            run.bold = True
            _set_run_font(run, font_name, font_size)
        elif part.startswith("$$") and part.endswith("$$"):
            # 行间公式
            inner = part[2:-2]
            converted = _convert_latex(inner)
            run = paragraph.add_run(converted)
            run.italic = True
            _set_run_font(run, font_name, font_size)
        elif part.startswith("$") and part.endswith("$"):
            # 行内公式
            inner = part[1:-1]
            converted = _convert_latex(inner)
            run = paragraph.add_run(converted)
            run.italic = True
            _set_run_font(run, font_name, font_size)
        else:
            run = paragraph.add_run(part)
            _set_run_font(run, font_name, font_size)


def _set_run_font(run, font_name: str, font_size: Pt):
    """设置 run 的字体。"""
    run.font.size = font_size
    run.font.name = font_name
    # 设置中文字体并清除主题字体（否则标题样式的主题字体会覆盖）
    rpr = run._element.get_or_add_rPr()
    rpr.attrib[qn("w:eastAsia")] = font_name
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is not None:
        for attr in ["w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"]:
            if qn(attr) in rFonts.attrib:
                del rFonts.attrib[qn(attr)]


def _set_cell_font(cell, text: str, bold: bool = False,
                    font_name: str = "Microsoft YaHei", font_size: Pt = Pt(10)):
    """设置单元格文本和字体。"""
    cell.text = ""
    p = cell.paragraphs[0]
    # 处理行内格式
    parts = re.split(r"(\*\*.*?\*\*|\$[^$]+\$)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
            _set_run_font(run, font_name, font_size)
        elif part.startswith("$") and part.endswith("$"):
            converted = _convert_latex(part[1:-1])
            run = p.add_run(converted)
            run.italic = True
            _set_run_font(run, font_name, font_size)
        else:
            run = p.add_run(part)
            run.bold = bold
            _set_run_font(run, font_name, font_size)


def _set_table_header_bg(table, color_hex: str = "dcecec"):
    """设置表头行背景色。"""
    for cell in table.rows[0].cells:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        shading = tcPr.makeelement(qn("w:shd"), {
            qn("w:fill"): color_hex,
            qn("w:val"): "clear",
        })
        tcPr.append(shading)


class MdToDocxConverter:
    """Markdown 转 Word 转换器。"""

    def __init__(self, font_name: str = "Microsoft YaHei", font_size: int = 11):
        self.font_name = font_name
        self.font_size = Pt(font_size)

    def convert(self, md_text: str, output_path: str | Path) -> Path:
        """将 Markdown 文本转换为 .docx 文件。

        Args:
            md_text: Markdown 文本内容
            output_path: 输出文件路径

        Returns:
            输出文件路径
        """
        output_path = Path(output_path)
        doc = Document()

        # 设置默认字体
        style = doc.styles["Normal"]
        style.font.name = self.font_name
        style.font.size = self.font_size
        style.element.rPr.rFonts.set(qn("w:eastAsia"), self.font_name)
        style.paragraph_format.line_spacing = 1.5

        # 设置标题样式字体
        for lvl in range(1, 7):
            style_name = f"Heading {lvl}"
            if style_name in doc.styles:
                h_style = doc.styles[style_name]
                h_style.font.name = self.font_name
                h_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
                h_rFonts = h_style.element.rPr.find(qn("w:rFonts"))
                if h_rFonts is not None:
                    h_rFonts.set(qn("w:eastAsia"), self.font_name)
                    for attr in ["w:asciiTheme", "w:hAnsiTheme", "w:eastAsiaTheme", "w:cstheme"]:
                        if qn(attr) in h_rFonts.attrib:
                            del h_rFonts.attrib[qn(attr)]

        # 设置页面边距
        for section in doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # 清理 LLM 输出的代码围栏
        md_text = self._strip_code_fences(md_text)

        # 按行解析
        lines = md_text.split("\n")
        i = 0
        while i < len(lines):
            line = lines[i]

            # 空行
            if not line.strip():
                i += 1
                continue

            # 水平分隔线
            if re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", line.strip()):
                # 添加一个带下边框的空段落
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                pPr = p._element.get_or_add_pPr()
                pBdr = pPr.makeelement(qn("w:pBdr"), {})
                bottom = pBdr.makeelement(qn("w:bottom"), {
                    qn("w:val"): "single",
                    qn("w:sz"): "6",
                    qn("w:space"): "1",
                    qn("w:color"): "999999",
                })
                pBdr.append(bottom)
                pPr.append(pBdr)
                i += 1
                continue

            # 标题
            header_match = re.match(r"^(#{1,6})\s+(.+)$", line)
            if header_match:
                level = len(header_match.group(1))
                text = header_match.group(2).strip()
                heading = doc.add_heading("", level=level)
                _parse_inline(heading, text, self.font_name, self._heading_size(level))
                i += 1
                continue

            # 表格
            if line.strip().startswith("|") and "|" in line.strip()[1:]:
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._convert_table(doc, table_lines)
                continue

            # 列表项
            list_match = re.match(r"^(\s*)-\s+(.+)$", line)
            if list_match:
                text = list_match.group(2)
                p = doc.add_paragraph(style="List Bullet")
                _parse_inline(p, text, self.font_name, self.font_size)
                i += 1
                continue

            # 缩进列表（如  - 子项）
            indent_match = re.match(r"^(\s{2,})-\s+(.+)$", line)
            if indent_match:
                text = indent_match.group(2)
                p = doc.add_paragraph(style="List Bullet 2")
                _parse_inline(p, text, self.font_name, self.font_size)
                i += 1
                continue

            # 普通段落
            # 收集连续的非空、非特殊行作为段落
            para_lines = [line]
            i += 1
            while i < len(lines):
                next_line = lines[i]
                if (not next_line.strip() or
                    next_line.strip().startswith("#") or
                    next_line.strip().startswith("|") or
                    re.match(r"^-{3,}$|^\*{3,}$|^_{3,}$", next_line.strip()) or
                    re.match(r"^(\s*)-\s+", next_line)):
                    break
                para_lines.append(next_line)
                i += 1
            text = " ".join(para_lines).strip()
            if text:
                p = doc.add_paragraph()
                _parse_inline(p, text, self.font_name, self.font_size)

        doc.save(str(output_path))
        return output_path

    def _heading_size(self, level: int) -> Pt:
        """返回标题字号。"""
        sizes = {1: Pt(18), 2: Pt(15), 3: Pt(13), 4: Pt(12), 5: Pt(11), 6: Pt(10)}
        return sizes.get(level, Pt(11))

    def _strip_code_fences(self, text: str) -> str:
        """移除 LLM 输出的代码围栏。"""
        text = text.strip()
        if text.startswith("```markdown"):
            text = text[len("```markdown"):]
        elif text.startswith("```"):
            text = text[3:]
        if text.endswith("```"):
            text = text[:-3]
        return text.strip()

    def _convert_table(self, doc, lines: list[str]):
        """将 Markdown 表格转换为 docx 表格。"""
        # 解析所有行
        rows_data = []
        for line in lines:
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            rows_data.append(cells)

        if len(rows_data) < 2:
            return

        # 过滤分隔行（如 |---|---|）
        data_rows = []
        for row in rows_data:
            if all(re.match(r"^[-:]+$", cell) for cell in row if cell):
                continue
            data_rows.append(row)

        if not data_rows:
            return

        # 确定列数
        max_cols = max(len(row) for row in data_rows)

        # 创建表格
        table = doc.add_table(rows=len(data_rows), cols=max_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 填充表格
        for row_idx, row_data in enumerate(data_rows):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx >= max_cols:
                    break
                cell = table.cell(row_idx, col_idx)
                is_header = (row_idx == 0)
                _set_cell_font(cell, cell_text, bold=is_header,
                              font_name=self.font_name, font_size=Pt(10))

        # 设置表头背景色
        if len(data_rows) > 0:
            _set_table_header_bg(table)

        # 表格后添加空段落
        doc.add_paragraph()
